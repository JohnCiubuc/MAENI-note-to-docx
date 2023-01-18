#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Fri Aug  5 21:06:31 2022

@author: John Ciubuc
"""
import math
import json
import sys
import base64
from docx import Document 
from docx.shared import RGBColor

def noteSection(note,item, subItem = -1):
    item = str(item)
    if item in note:
        data_item = note[item]   
        if subItem <= -1:
            if type(data_item) == str:
                return data_item.strip()
            elif type(data_item) == list:
                return data_item[0]
        else:
            subItem_s = str(subItem)
            if subItem_s in data_item:
                return data_item[subItem_s]
            elif subItem < len(data_item):
                if type(data_item) == list:
                    return data_item[subItem]
    return ""

def exportStudentNotes(json_text_string):
    json_text_string = base64.b64decode(json_text_string)
    jump = json.loads(json_text_string)
    note_db = jump['content']
    id = 0
    store = ''
    for i,idic in enumerate(note_db):
        if i != 0:
            store.update(note_db[idic])
        else:
            store = note_db[idic]
    note = store
    # new_store = store.po
    # Add a heading of level 0 (Also called Title)
    document = Document()
    # Choosing the top most section of the page
    section = document.sections[0]
     
    # Selecting the header
    header = section.header
     
    # Selecting the paragraph already present in
    # the header section
    header_para = header.paragraphs[0]
    
    header_para = header_para.add_run("\tThis document was exported from MAENI and may contain sensitive medical information.")
     
    # Adding the centred zoned header
    # header_para.text = 
    font = header_para.font
    font.color.rgb = RGBColor.from_string('8e0000')
    document.add_heading(f'Student Note ID: {id}', 0)
    document.add_heading('Chief Complaint', 1)
    document.add_paragraph(noteSection(note,20))
    document.add_heading('History of Presenting Illness', 1)
    document.add_paragraph(noteSection(note,21))
    document.add_heading('Review of Systems', 1)
    document.add_paragraph(noteSection(note,22))
    document.add_heading('History', 1)
    document.add_heading('Past Medical History', 2)
    document.add_paragraph(noteSection(note,23))
    document.add_heading('Past Surgical History', 2)
    document.add_paragraph(noteSection(note,24))
    document.add_heading('Medications', 2)
    document.add_paragraph(noteSection(note,25))
    document.add_heading('Allergies', 2)
    document.add_paragraph(noteSection(note,26))
    document.add_heading('Family History', 2)
    document.add_paragraph(noteSection(note,27))
    document.add_heading('Social History', 2)
    document.add_paragraph(noteSection(note,28))
    document.add_heading('Physical Exam', 1)
    document.add_heading('Vitals', 2)
    
    doc = f"""Heart Rate: {noteSection(note,29,0)}, Blood Pressure: {noteSection(note,29,1)}
Respiratory Rate: {noteSection(note,29,2)},  O2 Sat: {noteSection(note,29,3)}
Weight: {noteSection(note,29,4)}, Height: {noteSection(note,29,5)}"""
    
    document.add_paragraph(doc)
    document.add_heading('Exam', 2)
    document.add_paragraph(noteSection(note,39))
    document.add_heading('Data', 1)
    document.add_paragraph(noteSection(note,32))
    document.add_heading('Assessment and Plan', 1)
    document.add_heading('Summary Statement', 2)
    
    doc = f"""This is a {noteSection(note,38,1)} year old {noteSection(note,38,3)}, who is presenting today for {noteSection(note,38,6)}
The patient has a pertinent history of {noteSection(note,38,9)}
Patient's exam is remarkable for {noteSection(note,38,12)}
Patient's data is remarkable for {noteSection(note,38,15)}"""

    document.add_paragraph(doc)
    
    problem_list =  note['36'] 
    for problem_id in range(0,math.ceil(len(problem_list)/4)):
            document.add_heading(f'Problem {problem_id+1}:', 3)
            document.add_paragraph(noteSection(problem_list,problem_id*4+0))
            document.add_heading('Differential DX:', 3)
            document.add_paragraph(noteSection(problem_list,problem_id*4+1))
            document.add_heading('Diagnostic Plan:', 3)
            document.add_paragraph(noteSection(problem_list,problem_id*4+2))
            document.add_heading('Treatment Plan:', 3)
            document.add_paragraph(noteSection(problem_list,problem_id*4+3))
    
    document.save(f'.exports/Student Note {id}.docx')
    
# f = open('./test.txt', 'w')
# f.write(sys.argv[0])
# f.write("try arg")
# f.write(sys.argv[1])

if len(sys.argv) > 1:
    exportStudentNotes(sys.argv[1])
# f.write('wrote document')
# f.close()

print('exported document')